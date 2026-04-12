"""Construit les fixtures golden a partir du dossier 2. CR avec audio.

Extrait les 5 CR de reference (`Compte rendus/*.docx`) vers des fichiers
texte deterministes (`backend/tests/golden/fixtures/expected/*.txt`)
utilises par la regression v4 pour comparer la sortie du nouveau pipeline
avec le CR attendu.

L'audio brut n'est pas transcrit par ce script : la transcription est
faite on-demand par ``run_golden_v3.py`` (qui lui coute un appel Voxtral
par fichier). Ce script-ci est rapide, offline, idempotent.

Execution :
    python scripts/build_golden_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

REPO_ROOT: Path = Path(__file__).resolve().parent.parent
AUDIO_CR_DIR: Path = (
    REPO_ROOT / "Ressources" / "2. CR avec audio" / "Compte rendus"
)
FIXTURES_DIR: Path = (
    REPO_ROOT / "backend" / "tests" / "golden" / "fixtures"
)
EXPECTED_DIR: Path = FIXTURES_DIR / "expected"


def extract_docx_text(docx_path: Path) -> str:
    """Extrait le texte brut d'un fichier DOCX en preservant les lignes.

    Chaque paragraphe non vide devient une ligne. Les tableaux sont
    serialises ligne par ligne (une ligne = cellules separees par tab).
    """
    document = Document(str(docx_path))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text: str = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells: list[str] = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(cells))

    return "\n".join(lines) + "\n"


def write_fixture(name: str, content: str, target_dir: Path) -> Path:
    """Ecrit un fichier fixture avec creation des parents si besoin."""
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path: Path = target_dir / name
    output_path.write_text(content, encoding="utf-8")
    return output_path


def build_all_expected_fixtures() -> list[Path]:
    """Iterate les CR attendus et produit les fichiers texte correspondants."""
    produced: list[Path] = []

    if not AUDIO_CR_DIR.exists():
        raise FileNotFoundError(
            f"Dossier de reference introuvable : {AUDIO_CR_DIR}"
        )

    for docx_path in sorted(AUDIO_CR_DIR.glob("*.docx")):
        text: str = extract_docx_text(docx_path)
        fixture_name: str = f"{docx_path.stem}.txt"
        written: Path = write_fixture(fixture_name, text, EXPECTED_DIR)
        produced.append(written)

    return produced


def main() -> None:
    """Point d'entree CLI."""
    produced: list[Path] = build_all_expected_fixtures()
    print(f"[build_golden_fixtures] {len(produced)} fixtures generees")
    for path in produced:
        print(f"  -> {path.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
