"""Export de comptes-rendus anatomopathologiques au format Word (.docx).

Reproduit fidelement le format des CR anapath de reference :
- Police : Arial 10pt
- Titre : centre, gras, souligne, majuscules
- Renseignements cliniques : italique, justifie
- Sous-titres numerotes : gras + souligne, justifie
- Labels de section (Macroscopie, L'etude histologique) : gras
- Inclusion : italique
- CONCLUSION : gras + souligne + italique pour le label, gras pour le texte
- Corps : normal, justifie
- Tableaux IHC : grille 3 colonnes
- Marqueurs [A COMPLETER: xxx] : rouge gras

Architecture : chaque fonction fait UNE action.
"""

import io
import re
from typing import Any

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Pt, RGBColor, Cm  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]

_RED: RGBColor = RGBColor(220, 38, 38)
_BLACK: RGBColor = RGBColor(0, 0, 0)
_FONT_NAME: str = "Arial"
_FONT_SIZE: Pt = Pt(10)

_PATTERN_A_COMPLETER: re.Pattern[str] = re.compile(
    r"(\[A COMPLETER\s*:\s*[^\]]+\])", re.IGNORECASE
)

_PATTERN_BOLD_UNDERLINE: re.Pattern[str] = re.compile(
    r"\*\*__(.+?)__\*\*|__(.+?)__"
)

_PATTERN_INLINE_MARKUP: re.Pattern[str] = re.compile(
    r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)"
)


# ---------------------------------------------------------------------------
# Fonctions de formatage Word
# ---------------------------------------------------------------------------


def _set_run_font(run: object, bold: bool = False, italic: bool = False,
                  underline: bool = False, color: RGBColor = _BLACK) -> None:
    """Applique les proprietes de police a un run Word."""
    run.bold = bold  # type: ignore[attr-defined]
    run.italic = italic  # type: ignore[attr-defined]
    run.underline = underline  # type: ignore[attr-defined]
    run.font.name = _FONT_NAME  # type: ignore[attr-defined]
    run.font.size = _FONT_SIZE  # type: ignore[attr-defined]
    run.font.color.rgb = color  # type: ignore[attr-defined]


def _add_empty_paragraph(doc: Any) -> None:
    """Ajoute un paragraphe vide pour l'espacement."""
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE


def _add_title_paragraph(doc: Any, text: str,
                         is_conclusion: bool) -> None:
    """Ajoute un paragraphe titre (gras + souligne)."""
    is_main_title: bool = (
        text == text.upper()
        and not text.startswith(("1)", "2)", "3)", "4)", "5)"))
    )

    p = doc.add_paragraph()
    if is_main_title and not is_conclusion:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(text)
    _set_run_font(run, bold=True, italic=is_conclusion, underline=True)


def _add_header_paragraph(doc: Any, text: str, level: int) -> None:
    """Ajoute un paragraphe header markdown (# ou ##)."""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        _set_run_font(run, bold=True, underline=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        _set_run_font(run, bold=True, underline=(level == 2))


def _add_a_completer_run(paragraph: object, text: str) -> None:
    """Ajoute un run rouge gras pour un marqueur [A COMPLETER]."""
    run = paragraph.add_run(text)  # type: ignore[union-attr]
    _set_run_font(run, bold=True, color=_RED)


def _add_inline_run(paragraph: object, text: str) -> None:
    """Ajoute un run avec gestion du gras/italique inline."""
    run = paragraph.add_run("")  # type: ignore[union-attr]
    run.font.name = _FONT_NAME  # type: ignore[attr-defined]
    run.font.size = _FONT_SIZE  # type: ignore[attr-defined]
    run.font.color.rgb = _BLACK  # type: ignore[attr-defined]

    if text.startswith("***") and text.endswith("***"):
        run.text = text[3:-3]  # type: ignore[attr-defined]
        run.bold = True  # type: ignore[attr-defined]
        run.italic = True  # type: ignore[attr-defined]
    elif text.startswith("**") and text.endswith("**"):
        run.text = text[2:-2]  # type: ignore[attr-defined]
        run.bold = True  # type: ignore[attr-defined]
    elif text.startswith("*") and text.endswith("*") and len(text) > 2:
        run.text = text[1:-1]  # type: ignore[attr-defined]
        run.italic = True  # type: ignore[attr-defined]
    else:
        run.text = text  # type: ignore[attr-defined]


def _add_rich_text(paragraph: object, text: str) -> None:
    """Parse le texte inline avec gestion du gras, italique, et [A COMPLETER]."""
    parts_completer: list[str] = _PATTERN_A_COMPLETER.split(text)

    for part_c in parts_completer:
        if not part_c:
            continue

        if _PATTERN_A_COMPLETER.fullmatch(part_c):
            _add_a_completer_run(paragraph, part_c)
            continue

        # Retirer les **__TEXT__** et __TEXT__ (deja traites au niveau ligne)
        cleaned: str = _PATTERN_BOLD_UNDERLINE.sub(
            lambda m: m.group(1) or m.group(2) or "", part_c
        )

        parts: list[str] = _PATTERN_INLINE_MARKUP.split(cleaned)
        for part in parts:
            if part:
                _add_inline_run(paragraph, part)


def _add_table(doc: Any, table_lines: list[str]) -> None:
    """Ajoute un tableau IHC au format anapath standard."""
    rows_data: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        cells: list[str] = [c.strip() for c in line.strip().strip("|").split("|")]
        if idx == 1 and all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows_data.append(cells)

    if len(rows_data) < 1:
        return

    n_cols: int = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row):
            if c_idx < n_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_rich_text(p, cell_text)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
                        run.font.name = _FONT_NAME
                        run.font.size = _FONT_SIZE


def _configure_document_style(doc: Any) -> None:
    """Configure le style par defaut du document Word."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = _FONT_NAME
    font.size = _FONT_SIZE
    # Espacement inter-paragraphe compact comme un vrai CR
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _is_table_line(line: str) -> bool:
    """Verifie si une ligne est une ligne de tableau markdown."""
    return line.strip().startswith("|")


def _is_bold_underline_line(stripped: str) -> re.Match[str] | None:
    """Detecte un titre **__TEXT__** ou __TEXT__."""
    return _PATTERN_BOLD_UNDERLINE.match(stripped)


def _is_header_line(stripped: str) -> tuple[bool, int]:
    """Detecte un header markdown et retourne (is_header, level)."""
    if stripped.startswith("### "):
        return True, 3
    if stripped.startswith("## "):
        return True, 2
    if stripped.startswith("# "):
        return True, 1
    return False, 0


# ---------------------------------------------------------------------------
# Fonction principale d'export
# ---------------------------------------------------------------------------


def markdown_to_docx(markdown_text: str, title: str) -> bytes:
    """Convertit un rapport Markdown en document Word au format anapath standard."""
    doc: Any = Document()
    _configure_document_style(doc)

    lines: list[str] = markdown_text.split("\n")
    i: int = 0

    while i < len(lines):
        line: str = lines[i]
        stripped: str = line.strip()

        if not stripped:
            _add_empty_paragraph(doc)
            i += 1
            continue

        # Tableau markdown
        if _is_table_line(stripped) and i + 1 < len(lines):
            table_lines: list[str] = []
            while i < len(lines) and _is_table_line(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Titre **__TEXT__** ou __TEXT__
        bu_match: re.Match[str] | None = _is_bold_underline_line(stripped)
        if bu_match:
            text: str = bu_match.group(1) or bu_match.group(2) or ""
            is_conclusion: bool = "CONCLUSION" in text.upper()
            _add_title_paragraph(doc, text, is_conclusion)

            rest: str = stripped[bu_match.end():].strip()
            if rest:
                p = doc.paragraphs[-1]
                run = p.add_run(f" {rest}")
                _set_run_font(run)
            i += 1
            continue

        # Headers markdown
        is_header, level = _is_header_line(stripped)
        if is_header:
            header_text: str = stripped[level + 1:].strip()
            _add_header_paragraph(doc, header_text, level)
            i += 1
            continue

        # Separateur ---
        if stripped == "---":
            i += 1
            continue

        # Paragraphe normal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_rich_text(p, stripped)
        i += 1

    buffer: io.BytesIO = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Decoupage en sections
# ---------------------------------------------------------------------------


_SECTION_KEYWORDS: dict[str, list[str]] = {
    "renseignements_cliniques": [
        "renseignements cliniques",
        "renseignement clinique",
    ],
    "macroscopie": [
        "macroscopie",
        "examen macroscopique",
    ],
    "microscopie": [
        "microscopie",
        "etude histologique",
        "histologie",
        "etude cytologique",
        "l'etude histologique",
    ],
    "ihc": [
        "immunomarquage",
        "immunohistochimie",
    ],
    "biologie_moleculaire": [
        "biologie moleculaire",
        "biologie moléculaire",
    ],
}


def _detect_section_from_line(line: str) -> str | None:
    """Detecte si une ligne correspond au debut d'une section.

    Gere les multi-prelevements : une ligne comme
    ``**__2) Canal anal lateral gauche :__**`` cree une section
    ``prelevement_2`` pour eviter qu'elle soit avalee par la
    section precedente (ihc, microscopie, etc.).
    """
    clean_line: str = re.sub(r"[*_#|]", "", line).strip().lower()

    if re.search(r"\bconclusion\b", clean_line):
        return "conclusion"

    # Detection des numeros de prelevement (1), 2), 3)...)
    prel_match = re.match(r"^(\d+)\s*[)\.]\s*", clean_line)
    if prel_match:
        numero: str = prel_match.group(1)
        return f"prelevement_{numero}"

    for section_name, keywords in _SECTION_KEYWORDS.items():
        for keyword in keywords:
            if keyword in clean_line:
                return section_name

    return None


def split_report_sections(markdown_text: str) -> dict[str, str]:
    """Decoupe un rapport Markdown en sections nommees."""
    lines: list[str] = markdown_text.split("\n")
    sections: dict[str, list[str]] = {}
    current_section: str = "titre"
    sections[current_section] = []

    for line in lines:
        stripped: str = line.strip()

        if not stripped:
            if current_section in sections:
                sections[current_section].append(line)
            continue

        detected_section: str | None = _detect_section_from_line(stripped)

        if detected_section and detected_section != current_section:
            current_section = detected_section
            if current_section not in sections:
                sections[current_section] = []

        if current_section not in sections:
            sections[current_section] = []
        sections[current_section].append(line)

    result: dict[str, str] = {}
    for section_name, section_lines in sections.items():
        content: str = "\n".join(section_lines).strip()
        if content:
            result[section_name] = content

    return result
