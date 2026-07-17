"""Tests de l'export DOCX (fidélité de mise en page)."""

import io

from docx import Document

from export_docx import markdown_to_docx


def _render(cr: str) -> Document:
    return Document(io.BytesIO(markdown_to_docx(cr, "Test")))


def _find(doc: Document, needle: str):
    for p in doc.paragraphs:
        if needle.lower() in p.text.lower():
            return p
    return None


def test_title_bold_underline_uppercase():
    doc = _render("**__BIOPSIE PULMONAIRE__**\n**Macroscopie :**\nx")
    p = _find(doc, "BIOPSIE PULMONAIRE")
    assert p is not None
    assert any(r.bold for r in p.runs)
    assert any(r.underline for r in p.runs)


def test_title_inverted_markers_order():
    # __**X**__ (ordre inversé produit parfois par l'itération) doit aussi marcher.
    doc = _render("__**BIOPSIE DU SEIN**__\n**Microscopie :**\nx")
    p = _find(doc, "BIOPSIE DU SEIN")
    assert p is not None
    assert any(r.bold for r in p.runs) and any(r.underline for r in p.runs)
    # pas de marqueur résiduel
    assert "*" not in p.text and "_" not in p.text


def test_no_residual_markers():
    doc = _render(
        "**__TITRE__**\n**Macroscopie :**\nUne carotte.\n"
        "**__CONCLUSION :__**\n**Diagnostic final.**"
    )
    alltext = " ".join(p.text for p in doc.paragraphs)
    assert "**" not in alltext
    assert "__" not in alltext


def test_conclusion_bold():
    doc = _render("**__CONCLUSION :__**\n**Adenocarcinome infiltrant.**")
    p = _find(doc, "Adenocarcinome infiltrant")
    assert p is not None
    assert any(r.bold for r in p.runs)


def test_ihc_table_rendered():
    cr = (
        "**Microscopie :**\nx\n*Immunomarquage :*\n"
        "| Anticorps | Resultats |\n| TTF-1 | Positif |\n| PD-L1 | 5% |\n"
        "**__CONCLUSION :__**\n**x**"
    )
    doc = _render(cr)
    assert len(doc.tables) >= 1
    assert doc.tables[0].rows[0].cells[0].text.strip() == "Anticorps"


def test_section_extemporane_reconnue():
    """L'examen extemporane (26% des modeles CR du praticien, section CI-SIS)
    est decoupe dans sa propre section, avant la macroscopie."""
    from export_docx import split_report_sections

    cr = (
        "**__TITRE__**\n\n"
        "**Examen extemporané :**\n- Ganglion n°1 : sain.\n\n"
        "**Macroscopie :**\nPièce de lobectomie.\n\n"
        "**__CONCLUSION :__**\n**Adénocarcinome.**"
    )
    sections = split_report_sections(cr)
    assert "extemporane" in sections
    assert "Ganglion" in sections["extemporane"]
    # ordre : l'extemporane precede la macroscopie
    keys = list(sections)
    assert keys.index("extemporane") < keys.index("macroscopie")
