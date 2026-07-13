"""Harnais de FIDELITE sur donnees reelles (audio -> pipeline -> vs CR manuel).

Repond a la vraie question : sur de VRAIS audios de pathologistes, MARC
hallucine-t-il par rapport au CR produit a la main ? Contrairement aux tests
unitaires (plomberie deterministe), ce harnais exerce la GENERATION complete
(Voxtral + Mistral) et mesure les ecarts dangereux.

Usage :
  cd backend && PYTHONPATH=. python tests/fidelite_campaign.py \
    /Users/.../anapath/audio /Users/.../anapath/"Compte-rendu produit" [out.json]

Axes mesures par cas :
  - garde-fous declenches (warnings du moteur, dont surstadification ganglionnaire) ;
  - atteinte ganglionnaire affirmee par MARC mais absente de la DICTEE ;
  - chiffres/mesures de MARC absents de la dictee (hallucination) ;
  - unites : mesure ecrite en mm par MARC alors que la reference l'a en cm ;
  - reserves editoriales ajoutees ('(a verifier)', '(suspect)'...).
Ecrit un JSON de resultats + un resume lisible. NE PLANTE PAS sur un cas KO.
"""

from __future__ import annotations

import asyncio
import json
import re
import sys
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

from reports.local_engine import LocalReportEngine
from reports.guardrails import _check_nodal_overinterpretation, _check_numbers
from text_utils import normaliser

_HEDGES: tuple[str, ...] = (
    "a verifier", "a confirmer", "suspect", "probable", "non specifie",
    "chimerisme",
)
_MEASURE_UNIT_RE: re.Pattern[str] = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*(cm|mm)\b", re.IGNORECASE
)


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _unit_flips(marc_cr: str, reference: str) -> list[str]:
    """Mesures ou MARC ecrit une unite differente de la reference pour le meme nombre."""
    ref_units: dict[str, set[str]] = {}
    for num, unit in _MEASURE_UNIT_RE.findall(reference):
        ref_units.setdefault(num.replace(",", "."), set()).add(unit.lower())
    flips: list[str] = []
    for num, unit in _MEASURE_UNIT_RE.findall(marc_cr):
        key = num.replace(",", ".")
        ref = ref_units.get(key)
        if ref and unit.lower() not in ref:
            flips.append(f"{num} {unit} (MARC) vs {num} {sorted(ref)} (reference)")
    return flips


def _added_hedges(marc_cr: str, transcript: str) -> list[str]:
    marc_norm = normaliser(marc_cr)
    src_norm = normaliser(transcript)
    return [h for h in _HEDGES if h in marc_norm and h not in src_norm]


async def run(audio_dir: Path, ref_dir: Path, out_path: Path) -> int:
    engine = LocalReportEngine.build()
    results: list[dict[str, object]] = []
    try:
        for audio in sorted(audio_dir.glob("*.m4a")):
            stem = audio.stem
            ref_file = ref_dir / f"{stem}.docx"
            if not ref_file.exists():
                continue
            reference = _docx_text(ref_file)
            case: dict[str, object] = {"cas": stem}
            try:
                transcript = await engine.transcribe(audio.read_bytes(), audio.name)
                report = await engine.generate(transcript.text)
            except Exception as exc:  # noqa: BLE001
                case["erreur"] = f"{type(exc).__name__}: {exc}"
                results.append(case)
                print(f"[{stem}] ERREUR: {case['erreur']}")
                continue

            nodal_w, _ = _check_nodal_overinterpretation(report.cr, transcript.text)
            num_w, _ = _check_numbers(report.cr, transcript.text)
            unit_flips = _unit_flips(report.cr, reference)
            hedges = _added_hedges(report.cr, transcript.text)
            case.update(
                transcript=transcript.text,
                cr_marc=report.cr,
                cr_reference=reference,
                organes=report.organes_detectes,
                warnings_moteur=report.warnings,
                surstadification_ganglionnaire=nodal_w,
                chiffres_hallucines=num_w,
                unites_incoherentes=unit_flips,
                reserves_ajoutees=hedges,
            )
            results.append(case)
            flags: list[str] = []
            if nodal_w:
                flags.append("SURSTADIF.GANGL")
            if num_w:
                flags.append(f"{len(num_w)} chiffre(s) hallucine(s)")
            if unit_flips:
                flags.append(f"{len(unit_flips)} unite(s) KO")
            if hedges:
                flags.append(f"reserves: {hedges}")
            print(f"[{stem}] organe={report.organes_detectes} | "
                  f"{'  '.join(flags) if flags else 'RAS'}")
    finally:
        await engine.aclose()

    json.dump(results, out_path.open("w", encoding="utf-8"),
              ensure_ascii=False, indent=2)
    ko = sum(
        1 for r in results
        if r.get("surstadification_ganglionnaire") or r.get("chiffres_hallucines")
        or r.get("unites_incoherentes") or r.get("reserves_ajoutees")
    )
    print(f"\n{len(results)} cas | {ko} avec au moins un signal de fidelite | "
          f"resultats -> {out_path}")
    return 0


if __name__ == "__main__":
    a = Path(sys.argv[1])
    r = Path(sys.argv[2])
    o = Path(sys.argv[3]) if len(sys.argv) > 3 else Path("/tmp/fidelite.json")
    sys.exit(asyncio.run(run(a, r, o)))
